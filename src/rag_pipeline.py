"""
Lightweight RAG Pipeline Implementation
Fallback version that works without heavy ML dependencies
"""

import os
import logging
from typing import List, Dict, Optional, Any
from pathlib import Path
import tempfile
import json
from concurrent.futures import ThreadPoolExecutor
import asyncio

import ollama
import chromadb
from chromadb.config import Settings
import PyPDF2
from docx import Document
import numpy as np

# Optional heavy imports with fallbacks
try:
    from sentence_transformers import SentenceTransformer
    SENTENCE_TRANSFORMERS_AVAILABLE = True
except ImportError:
    SENTENCE_TRANSFORMERS_AVAILABLE = False
    logging.warning("sentence-transformers not available, using simple chunking")

try:
    from sklearn.metrics.pairwise import cosine_similarity as sklearn_cosine_similarity
    SKLEARN_AVAILABLE = True
except ImportError:
    SKLEARN_AVAILABLE = False
    logging.warning("scikit-learn not available, using numpy cosine similarity")

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


def simple_cosine_similarity(a: np.ndarray, b: np.ndarray) -> float:
    """Simple cosine similarity using numpy"""
    dot_product = np.dot(a, b)
    norm_a = np.linalg.norm(a)
    norm_b = np.linalg.norm(b)
    return dot_product / (norm_a * norm_b)


class DocumentProcessor:
    """Simple document processing using PyPDF2 and python-docx"""
    
    def __init__(self):
        pass
    
    def process_document(self, file_path: str) -> Dict[str, Any]:
        """Process document and extract content"""
        file_extension = Path(file_path).suffix.lower()
        
        if file_extension == '.pdf':
            return self.process_pdf(file_path)
        elif file_extension in ['.docx', '.doc']:
            return self.process_docx(file_path)
        elif file_extension in ['.txt', '.md']:
            return self.process_text(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_extension}")
    
    def process_pdf(self, pdf_path: str) -> Dict[str, Any]:
        """Process PDF and extract text content"""
        try:
            with open(pdf_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text_content = []
                
                for page in pdf_reader.pages:
                    text_content.append(page.extract_text())
                
                full_text = '\n'.join(text_content)
                
                content = {
                    'text': full_text,
                    'metadata': {
                        'source': pdf_path,
                        'num_pages': len(pdf_reader.pages),
                        'title': os.path.basename(pdf_path),
                        'file_type': 'pdf'
                    }
                }
                
                logger.info(f"Processed PDF: {pdf_path}, extracted {len(full_text)} characters from {len(pdf_reader.pages)} pages")
                return content
                
        except Exception as e:
            logger.error(f"Error processing PDF {pdf_path}: {str(e)}")
            raise
    
    def process_docx(self, docx_path: str) -> Dict[str, Any]:
        """Process DOCX file and extract text content"""
        try:
            doc = Document(docx_path)
            text_content = []
            
            for paragraph in doc.paragraphs:
                text_content.append(paragraph.text)
            
            full_text = '\n'.join(text_content)
            
            content = {
                'text': full_text,
                'metadata': {
                    'source': docx_path,
                    'num_paragraphs': len(doc.paragraphs),
                    'title': os.path.basename(docx_path),
                    'file_type': 'docx'
                }
            }
            
            logger.info(f"Processed DOCX: {docx_path}, extracted {len(full_text)} characters")
            return content
            
        except Exception as e:
            logger.error(f"Error processing DOCX {docx_path}: {str(e)}")
            raise
    
    def process_text(self, text_path: str) -> Dict[str, Any]:
        """Process text file"""
        try:
            with open(text_path, 'r', encoding='utf-8') as file:
                content_text = file.read()
            
            content = {
                'text': content_text,
                'metadata': {
                    'source': text_path,
                    'title': os.path.basename(text_path),
                    'file_type': 'text'
                }
            }
            
            logger.info(f"Processed text file: {text_path}, extracted {len(content_text)} characters")
            return content
            
        except Exception as e:
            logger.error(f"Error processing text file {text_path}: {str(e)}")
            raise


class OllamaEmbedder:
    """Embedding generation using Ollama"""
    
    def __init__(self, model_name: str = "nomic-embed-text"):
        self.model_name = model_name
        # Use environment variable for Ollama base URL
        base_url = os.getenv('OLLAMA_BASE_URL', 'http://localhost:11434')
        self.client = ollama.Client(host=base_url)
        self._ensure_model()
    
    def _ensure_model(self):
        """Ensure the embedding model is available"""
        try:
            models = self.client.list()
            model_names = [model['name'] for model in models['models']]
            
            if self.model_name not in model_names:
                logger.info(f"Pulling embedding model: {self.model_name}")
                self.client.pull(self.model_name)
        except Exception as e:
            logger.warning(f"Could not verify model availability: {str(e)}")
    
    def embed_text(self, text: str) -> List[float]:
        """Generate embeddings for text"""
        try:
            response = self.client.embeddings(
                model=self.model_name,
                prompt=text
            )
            return response['embedding']
        except Exception as e:
            logger.error(f"Error generating embedding: {str(e)}")
            raise
    
    def embed_batch(self, texts: List[str]) -> List[List[float]]:
        """Generate embeddings for multiple texts"""
        embeddings = []
        for text in texts:
            embeddings.append(self.embed_text(text))
        return embeddings


class SimpleChunker:
    """Simple text chunking without ML dependencies"""
    
    def __init__(self, chunk_size: int = 1000, overlap: int = 200):
        self.chunk_size = chunk_size
        self.overlap = overlap
    
    def chunk_text(self, text: str) -> List[str]:
        """Simple chunking by character count with overlap"""
        if len(text) <= self.chunk_size:
            return [text]
        
        chunks = []
        start = 0
        
        while start < len(text):
            end = start + self.chunk_size
            
            # Try to break at word boundary
            if end < len(text):
                # Find last space before the cut-off
                last_space = text.rfind(' ', start, end)
                if last_space > start:
                    end = last_space
            
            chunk = text[start:end].strip()
            if chunk:
                chunks.append(chunk)
            
            start = end - self.overlap
            if start >= len(text):
                break
        
        return chunks


class SemanticChunker:
    """Semantic-based text chunking with fallback to simple chunking"""
    
    def __init__(self, model_name: str = "all-MiniLM-L6-v2"):
        if SENTENCE_TRANSFORMERS_AVAILABLE:
            try:
                self.model = SentenceTransformer(model_name)
                self.use_semantic = True
                logger.info("Using semantic chunking with sentence-transformers")
            except Exception as e:
                logger.warning(f"Failed to load sentence-transformers model: {e}")
                self.use_semantic = False
        else:
            self.use_semantic = False
            
        if not self.use_semantic:
            self.simple_chunker = SimpleChunker()
            logger.info("Using simple chunking as fallback")
    
    def chunk_by_similarity(self, text: str, threshold: float = 0.7, 
                          max_chunk_size: int = 1000) -> List[str]:
        """Chunk text based on semantic similarity or fallback to simple chunking"""
        if not self.use_semantic:
            return self.simple_chunker.chunk_text(text)
        
        sentences = self._split_into_sentences(text)
        if len(sentences) <= 1:
            return [text]
        
        try:
            embeddings = self.model.encode(sentences)
            chunks = []
            current_chunk = [sentences[0]]
            current_size = len(sentences[0])
            
            for i in range(1, len(sentences)):
                if SKLEARN_AVAILABLE:
                    similarity = sklearn_cosine_similarity(
                        [embeddings[i-1]], 
                        [embeddings[i]]
                    )[0][0]
                else:
                    similarity = simple_cosine_similarity(
                        embeddings[i-1], 
                        embeddings[i]
                    )
                
                sentence_size = len(sentences[i])
                
                if (similarity > threshold and 
                    current_size + sentence_size < max_chunk_size):
                    current_chunk.append(sentences[i])
                    current_size += sentence_size
                else:
                    if current_chunk:
                        chunks.append(' '.join(current_chunk))
                    current_chunk = [sentences[i]]
                    current_size = sentence_size
            
            if current_chunk:
                chunks.append(' '.join(current_chunk))
            
            return chunks
        except Exception as e:
            logger.warning(f"Semantic chunking failed: {e}, falling back to simple chunking")
            return self.simple_chunker.chunk_text(text)
    
    def _split_into_sentences(self, text: str) -> List[str]:
        """Simple sentence splitting"""
        sentences = text.replace('\n', ' ').split('. ')
        return [s.strip() + '.' for s in sentences if s.strip()]


# Define the missing classes locally to avoid import issues


class VectorStore:
    """Vector storage using ChromaDB"""
    
    def __init__(self, collection_name: str = "documents", 
                 persist_directory: str = "./chroma_db"):
        self.persist_directory = persist_directory
        self.client = chromadb.PersistentClient(path=persist_directory)
        self.collection = self.client.get_or_create_collection(
            name=collection_name,
            metadata={"hnsw:space": "cosine"}
        )
    
    def add_documents(self, documents: List[Dict[str, Any]]):
        """Add documents to vector store"""
        if not documents:
            return
        
        ids = [doc['id'] for doc in documents]
        embeddings = [doc['embedding'] for doc in documents]
        metadatas = [doc['metadata'] for doc in documents]
        documents_text = [doc['text'] for doc in documents]
        
        try:
            self.collection.add(
                ids=ids,
                embeddings=embeddings,
                metadatas=metadatas,
                documents=documents_text
            )
            logger.info(f"Added {len(documents)} documents to vector store")
        except Exception as e:
            logger.error(f"Error adding documents to vector store: {str(e)}")
            raise
    
    def search(self, query_embedding: List[float], k: int = 5) -> List[Dict[str, Any]]:
        """Search for similar documents"""
        try:
            results = self.collection.query(
                query_embeddings=[query_embedding],
                n_results=k,
                include=['documents', 'metadatas', 'distances']
            )
            
            return [
                {
                    'text': doc,
                    'metadata': meta,
                    'score': 1 - dist  # Convert distance to similarity score
                }
                for doc, meta, dist in zip(
                    results['documents'][0],
                    results['metadatas'][0],
                    results['distances'][0]
                )
            ]
        except Exception as e:
            logger.error(f"Error searching vector store: {str(e)}")
            return []
    
    def get_collection_stats(self) -> Dict[str, Any]:
        """Get statistics about the collection"""
        try:
            count = self.collection.count()
            return {
                'document_count': count,
                'collection_name': self.collection.name
            }
        except Exception as e:
            logger.error(f"Error getting collection stats: {str(e)}")
            return {'document_count': 0}
    
    def get_documents(self) -> List[Dict[str, Any]]:
        """Get list of documents with metadata"""
        try:
            # Get all documents from the collection
            results = self.collection.get(include=['metadatas'])
            
            # Group by source document
            docs_by_source = {}
            for metadata in results['metadatas']:
                source = metadata.get('source', 'Unknown')
                title = metadata.get('title', 'Unknown Document')
                doc_type = metadata.get('document_type', 'unknown')
                
                if source not in docs_by_source:
                    docs_by_source[source] = {
                        'title': title,
                        'source': source,
                        'document_type': doc_type,
                        'chunk_count': 0
                    }
                docs_by_source[source]['chunk_count'] += 1
            
            return list(docs_by_source.values())
            
        except Exception as e:
            logger.error(f"Error getting documents: {str(e)}")
            return []
    
    def clear_documents(self) -> Dict[str, Any]:
        """Clear all documents from the vector store"""
        try:
            # Delete the current collection and recreate it
            self.client.delete_collection(self.collection.name)
            self.collection = self.client.get_or_create_collection(
                name=self.collection.name,
                metadata={"hnsw:space": "cosine"}
            )
            
            logger.info("Successfully cleared all documents from vector store")
            return {"message": "All documents cleared successfully", "status": "success"}
            
        except Exception as e:
            logger.error(f"Error clearing documents: {str(e)}")
            return {"message": f"Failed to clear documents: {str(e)}", "status": "error"}


class OllamaLLM:
    """Local LLM using Ollama"""
    
    def __init__(self, model_name: str = "llama3.2:1b"):
        self.model_name = model_name
        # Use environment variable for Ollama base URL
        base_url = os.getenv('OLLAMA_BASE_URL', 'http://localhost:11434')
        self.client = ollama.Client(host=base_url)
        self._ensure_model()
    
    def _ensure_model(self):
        """Ensure the LLM model is available"""
        try:
            models = self.client.list()
            model_names = [model['name'] for model in models['models']]
            
            if self.model_name not in model_names:
                logger.info(f"Pulling LLM model: {self.model_name}")
                self.client.pull(self.model_name)
        except Exception as e:
            logger.warning(f"Could not verify model availability: {str(e)}")
    
    def generate_response(self, prompt: str, context: str = "") -> str:
        """Generate response using local LLM"""
        try:
            full_prompt = f"""
Context: {context}

Question: {prompt}

Please provide a comprehensive answer based on the context provided. If the context doesn't contain enough information to answer the question, say so clearly.
"""
            
            response = self.client.chat(
                model=self.model_name,
                messages=[{"role": "user", "content": full_prompt}],
                stream=False
            )
            
            return response['message']['content']
            
        except Exception as e:
            logger.error(f"Error generating response: {str(e)}")
            return f"I apologize, but I encountered an error while generating a response: {str(e)}"


class BatchProcessor:
    """Batch processing for multiple documents"""
    
    def __init__(self, rag_pipeline: 'RAGPipeline', max_workers: int = 4):
        self.rag_pipeline = rag_pipeline
        self.max_workers = max_workers
    
    async def process_documents_batch(self, file_paths: List[str]) -> List[Dict[str, Any]]:
        """Process multiple documents in parallel"""
        with ThreadPoolExecutor(max_workers=self.max_workers) as executor:
            loop = asyncio.get_event_loop()
            tasks = [
                loop.run_in_executor(
                    executor, 
                    self.rag_pipeline.ingest_document, 
                    file_path
                )
                for file_path in file_paths
            ]
            results = await asyncio.gather(*tasks, return_exceptions=True)
        
        # Process results and handle exceptions
        processed_results = []
        for i, result in enumerate(results):
            if isinstance(result, Exception):
                processed_results.append({
                    'file': file_paths[i],
                    'status': 'error',
                    'error': str(result)
                })
            else:
                processed_results.append({
                    'file': file_paths[i],
                    'status': 'success',
                    'result': result
                })
        
        return processed_results


class RAGPipeline:
    """Main RAG pipeline with lightweight dependencies"""
    
    def __init__(self, 
                 embedding_model: str = "nomic-embed-text",
                 llm_model: str = "llama3.2:1b",
                 collection_name: str = "documents"):
        
        self.embedder = OllamaEmbedder(embedding_model)
        self.llm = OllamaLLM(llm_model)
        self.vector_store = VectorStore(collection_name)
        self.processor = DocumentProcessor()
        self.chunker = SemanticChunker()  # Will fall back to simple chunking if needed
        
        logger.info("RAG Pipeline initialized successfully")
    
    def ingest_document(self, pdf_path: str, use_semantic_chunking: bool = True, original_filename: str = None) -> Dict[str, Any]:
        """Ingest a document into the RAG system"""
        try:
            # Process document
            document = self.processor.process_document(pdf_path)
            text = document['text']
            metadata = document['metadata']
            
            # Update metadata with original filename if provided
            if original_filename:
                metadata['title'] = original_filename
                metadata['original_filename'] = original_filename
                metadata['source'] = original_filename
            
            if not text.strip():
                raise ValueError("No text content extracted from document")
            
            # Chunk text
            if use_semantic_chunking:
                chunks = self.chunker.chunk_by_similarity(text)
            else:
                simple_chunker = SimpleChunker()
                chunks = simple_chunker.chunk_text(text)
            
            # Generate embeddings
            embeddings = self.embedder.embed_batch(chunks)
            
            # Prepare documents for vector store
            documents = []
            for i, (chunk, embedding) in enumerate(zip(chunks, embeddings)):
                doc_metadata = metadata.copy()
                doc_metadata.update({
                    'chunk_id': i,
                    'total_chunks': len(chunks),
                    'chunk_text': chunk[:200] + "..." if len(chunk) > 200 else chunk
                })
                
                documents.append({
                    'id': f"{metadata['title']}_{i}",
                    'text': chunk,
                    'embedding': embedding,
                    'metadata': doc_metadata
                })
            
            # Store in vector database
            self.vector_store.add_documents(documents)
            
            result = {
                "message": "Document ingested successfully",
                "chunks_created": len(chunks),
                "source": metadata['title'],
                "metadata": metadata
            }
            
            logger.info(f"Successfully ingested document: {metadata['title']} with {len(chunks)} chunks")
            return result
            
        except Exception as e:
            logger.error(f"Error ingesting document: {str(e)}")
            raise
    
    def query(self, question: str, context_limit: int = 5) -> Dict[str, Any]:
        """Query the RAG system"""
        try:
            # Generate embedding for query
            query_embedding = self.embedder.embed_text(question)
            
            # Search vector database
            search_results = self.vector_store.search(query_embedding, k=context_limit)
            
            if not search_results:
                return {
                    "answer": "I don't have any relevant information to answer your question.",
                    "sources": [],
                    "context_used": 0,
                    "relevance_scores": []
                }
            
            # Build context
            context_chunks = []
            sources = []
            relevance_scores = []
            
            for result in search_results:
                context_chunks.append(result['text'])
                sources.append({
                    'source': result['metadata'].get('source', 'Unknown'),
                    'chunk_id': result['metadata'].get('chunk_id', 0),
                    'relevance_score': result.get('score', 0.0)
                })
                relevance_scores.append(result.get('score', 0.0))
            
            context = "\n\n".join(context_chunks)
            
            # Generate response
            answer = self.llm.generate_response(question, context)
            
            return {
                "answer": answer,
                "sources": sources,
                "context_used": len(context_chunks),
                "relevance_scores": relevance_scores
            }
            
        except Exception as e:
            logger.error(f"Error querying RAG system: {str(e)}")
            raise
    
    def get_stats(self) -> Dict[str, Any]:
        """Get pipeline statistics"""
        try:
            vector_stats = self.vector_store.get_collection_stats()
            
            return {
                "pipeline_status": "healthy",
                "vector_store_stats": vector_stats,
                "embedding_model": self.embedder.model_name,
                "llm_model": self.llm.model_name,
                "semantic_chunking_available": SENTENCE_TRANSFORMERS_AVAILABLE,
                "sklearn_available": SKLEARN_AVAILABLE
            }
        except Exception as e:
            logger.error(f"Error getting stats: {str(e)}")
            return {"pipeline_status": "error", "error": str(e)}
    
    def get_documents(self) -> List[Dict[str, Any]]:
        """Get list of ingested documents"""
        return self.vector_store.get_documents()
    
    def clear_documents(self) -> Dict[str, Any]:
        """Clear all documents from the vector store"""
        try:
            result = self.vector_store.clear_documents()
            logger.info("Cleared all documents from vector store")
            return result
        except Exception as e:
            logger.error(f"Error clearing documents: {str(e)}")
            raise 